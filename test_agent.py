"""Checks for the things that break silently.

No network, no LLM calls. Run with: .venv/bin/python -m pytest test_agent.py
or just: .venv/bin/python test_agent.py
"""

import asyncio
import io
import json
import os
import pathlib
import tempfile

from google.adk.memory.vertex_ai_memory_bank_service import VertexAiMemoryBankService
from google.adk.sessions.vertex_ai_session_service import VertexAiSessionService
from google.adk.tools.agent_tool import AgentTool
from google.adk.tools.base_tool import BaseTool
from google.adk.tools.function_tool import FunctionTool
from google.genai import types

from Job_Helper_agent.agent import SPECIALISTS, root_agent
from Job_Helper_agent.callbacks import require_real_user
from Job_Helper_agent.runtime import REQUIRED_ENV, build_runner
from Job_Helper_agent.tools import list_applications, track_application
from placement_agent.agent import root_agent as placement_root_agent
from placement_agent.resume_parser import parse_resume

# Built-in Gemini tools, which cannot share an agent with function tools.
BUILT_IN_NAMES = {"google_search", "url_context", "code_execution", "computer_use"}

EXPECTED_OUTPUT_KEYS = {
    "profile_agent": "profile",
    # SequentialAgent wrappers: the output_key lives on their sub-agents.
    "company_agent": None,
    "alumni_agent": None,
    "matching_agent": "matches",
    "verification_agent": None,
    "resume_gap_agent": "gaps",
    "outreach_agent": None,
    "tracker_agent": None,
    "coach_agent": None,
}


def _leaf_agents(agents):
    """Flatten SequentialAgent wrappers to the agents that actually hold tools.

    company_agent and alumni_agent are two-step pipelines (search, then
    structure), so the tool-safety checks must reach inside them.
    """
    out = []
    for agent in agents:
        subs = getattr(agent, "sub_agents", None)
        if subs:
            out.extend(_leaf_agents(subs))
        else:
            out.append(agent)
    return out


def _tool_name(tool) -> str:
    if isinstance(tool, AgentTool):
        return tool.agent.name
    return getattr(tool, "name", None) or getattr(tool, "__name__", repr(tool))


def _is_built_in(tool) -> bool:
    return _tool_name(tool) in BUILT_IN_NAMES


def _is_function_tool(tool) -> bool:
    # A plain python function passed to `tools=` is wrapped by ADK, so accept
    # both shapes.
    return callable(tool) and not isinstance(tool, BaseTool) or isinstance(
        tool, (FunctionTool, AgentTool)
    )


def test_root_has_all_specialists_plus_infrastructure_tools():
    names = [_tool_name(t) for t in root_agent.tools]
    for agent in SPECIALISTS:
        assert agent.name in names, f"{agent.name} not wired into root"

    # These two are easy to drop in a refactor and their absence is silent:
    # without load_artifacts the agent can never read an uploaded resume and
    # will likely invent one; without preload_memory a returning student's
    # application history never comes back.
    assert "load_artifacts" in names, "root cannot read uploaded resumes"
    assert "preload_memory" in names, "root cannot recall application history"
    assert "alumni_search_links" in names, "no fallback when alumni search is empty"
    assert len(root_agent.tools) == len(SPECIALISTS) + 3


def test_output_keys_match_the_spec():
    for agent in SPECIALISTS:
        expected = EXPECTED_OUTPUT_KEYS[agent.name]
        if getattr(agent, "sub_agents", None):
            # A SequentialAgent has no output_key of its own; its sub-agents
            # carry them. Checked separately below.
            assert getattr(agent, "output_key", None) is None
            continue
        assert agent.output_key == expected, (
            f"{agent.name} output_key is {agent.output_key!r}, expected"
            f" {expected!r} -- downstream agents read this via {{key?}}"
        )


def test_sequential_specialists_keep_their_state_contract():
    """The wrappers hide output_key, so assert it on the halves that have it.

    The search half still writes the prose key downstream agents read via
    {key?}; the structure half writes the *_data key the A2UI renderer draws.
    """
    from Job_Helper_agent.agent import (
        alumni_search_agent,
        alumni_structure_agent,
        company_search_agent,
        company_structure_agent,
    )

    assert company_search_agent.output_key == "companies"
    assert company_structure_agent.output_key == "companies_data"
    assert alumni_search_agent.output_key == "alumni"
    assert alumni_structure_agent.output_key == "alumni_data"

    # Controlled generation and the Search tool cannot coexist -- Gemini
    # returns 400 INVALID_ARGUMENT. That is why structuring is a second agent
    # rather than an output_schema on the search agent, and why the search
    # halves must never acquire one.
    assert company_search_agent.output_schema is None
    assert alumni_search_agent.output_schema is None
    assert not (company_structure_agent.tools or [])
    assert not (alumni_structure_agent.tools or [])


def test_no_agent_can_fetch_linkedin():
    """url_context would fetch any URL a student pasted, including LinkedIn.

    robots.txt is `User-agent: * / Disallow: /`, so that would be prohibited
    automated access originating from our production agent. Job descriptions
    now go through fetch_job_description, which blocklists such domains in
    code -- instructions alone have already failed us twice.
    """
    for agent in [root_agent, *_leaf_agents(SPECIALISTS)]:
        assert "url_context" not in [_tool_name(t) for t in (agent.tools or [])], (
            f"{agent.name} holds url_context, which will fetch any URL given to it"
        )


def test_no_agent_mixes_builtin_and_function_tools():
    """The constraint ADK will not enforce for us.

    A Gemini built-in cannot share an LlmAgent with custom function tools. ADK
    raises nothing (llm_agent.py:139-176 silently rewrites instead), and the
    auto-wrap workaround defaults to off, so this fails at the Gemini API --
    potentially only once deployed.
    """
    for agent in [root_agent, *_leaf_agents(SPECIALISTS)]:
        tools = list(getattr(agent, "tools", None) or [])
        built_ins = [t for t in tools if _is_built_in(t)]
        if not built_ins:
            continue
        assert len(tools) == 1, (
            f"{agent.name} holds built-in {_tool_name(built_ins[0])!r} alongside"
            f" {[_tool_name(t) for t in tools if not _is_built_in(t)]}."
            " Gemini rejects this at request time."
        )


class _FakeCallbackContext:
    def __init__(self, user_id):
        self.user_id = user_id


def test_identity_guard_rejects_only_the_shared_bucket():
    """Reject values that put two students in ONE bucket. Nothing else.

    `default-user-id` is Agent Engine's silent fallback and is identical for
    every caller -- one bucket, everyone's data in it. That is a leak and stays
    refused.

    `A2A_USER_{context_id}` is ADK's A2A fallback. Measured 2026-07-28: Gemini
    Enterprise forwards no end-user identity to an A2A agent at all -- no email
    header, nothing in message metadata, and the only credential present is the
    Discovery Engine service agent, which is the same for everyone. The
    context_id is a per-conversation UUID, so accepting it FRAGMENTS the scope
    into one bucket per conversation rather than collapsing it. Forgetful, not
    leaky -- and refusing it instead simply took the agent offline.
    """
    # One shared bucket for everyone. Still refused.
    assert require_real_user(_FakeCallbackContext("default-user-id")) is not None
    assert require_real_user(_FakeCallbackContext("")) is not None
    assert require_real_user(_FakeCallbackContext(None)) is not None

    # No context id means no conversation to scope to.
    assert require_real_user(_FakeCallbackContext("A2A_USER_")) is not None

    # Conversation-scoped: private per chat, so allowed.
    assert require_real_user(_FakeCallbackContext("A2A_USER_ctx-abc123")) is None

    # A real signed-in student, if identity ever becomes available.
    assert require_real_user(_FakeCallbackContext("student@example.com")) is None


def test_identity_headers_exclude_client_assertable_names():
    """Only proxy-managed headers may name a user.

    `x-user-email` is not managed by any Google proxy, so nothing strips a
    client-supplied copy -- anyone able to reach the service could assert
    another student's identity with it.
    """
    from Job_Helper_agent.identity import IDENTITY_HEADERS

    assert "x-user-email" not in IDENTITY_HEADERS
    assert all(h.startswith("x-goog-") for h in IDENTITY_HEADERS), IDENTITY_HEADERS


class _FakeContext:
    """Minimal stand-in for ToolContext: the tools only touch .state."""

    def __init__(self):
        self.state = {}


def test_track_application_appends_then_updates():
    ctx = _FakeContext()

    first = track_application("Google", "SWE Intern", "Applied", "via careers page", ctx)
    assert first["total"] == 1
    assert ctx.state["applications"][0]["company"] == "Google"

    track_application("Atlassian", "Backend Intern", "Applied", "", ctx)
    listed = list_applications(ctx)
    assert listed["total"] == 2

    # Same company + role must update in place, not duplicate -- otherwise the
    # coach agent double-counts the pipeline.
    moved = track_application("google", "swe intern", "Interview", "onsite Aug 3", ctx)
    assert "updated" in moved
    assert moved["total"] == 2
    assert ctx.state["applications"][0]["status"] == "Interview"
    assert ctx.state["applications"][0]["notes"] == "onsite Aug 3"


def test_track_application_rejects_unknown_status():
    ctx = _FakeContext()
    result = track_application("Google", "SWE", "Ghosted", "", ctx)
    assert "error" in result
    assert ctx.state.get("applications") in (None, [])


def test_build_runner_refuses_to_start_without_backing_store_config():
    # Fail loudly at boot rather than silently serving from memory: an
    # in-memory session service on Cloud Run loses a student's tracked
    # applications the moment the instance recycles.
    saved = {k: os.environ.pop(k, None) for k in REQUIRED_ENV}
    try:
        raised = None
        try:
            build_runner()
        except RuntimeError as e:
            raised = e
        assert raised is not None, "build_runner started with no configuration"
        assert "AGENT_ENGINE_ID" in str(raised)

        # Partial config is the likelier deploy mistake than none at all, and
        # the message has to name everything still missing or the operator
        # fixes one variable per redeploy.
        os.environ["GOOGLE_CLOUD_PROJECT"] = "test-project"
        raised = None
        try:
            build_runner()
        except RuntimeError as e:
            raised = e
        assert raised is not None, "build_runner started with partial configuration"
        assert "GOOGLE_CLOUD_LOCATION" in str(raised)
        assert "AGENT_ENGINE_ID" in str(raised)
        assert "GOOGLE_CLOUD_PROJECT" not in str(raised)
    finally:
        for k in REQUIRED_ENV:
            os.environ.pop(k, None)
        for k, v in saved.items():
            if v is not None:
                os.environ[k] = v


def test_build_runner_uses_persistent_services_not_in_memory_defaults():
    saved = {k: os.environ.get(k) for k in REQUIRED_ENV}
    os.environ["GOOGLE_CLOUD_PROJECT"] = "test-project"
    os.environ["GOOGLE_CLOUD_LOCATION"] = "us-central1"
    os.environ["AGENT_ENGINE_ID"] = "1234567890"
    try:
        runner = build_runner()
        # These three are the regression guard for the migration. Each default
        # is a silent data-loss path, not a performance nicety.
        assert isinstance(runner.session_service, VertexAiSessionService)
        assert isinstance(runner.memory_service, VertexAiMemoryBankService)
        assert runner.agent is not None
        # app_name is the Memory Bank retrieval scope. The Agent Engine
        # template this replaces scoped memories to the engine id, so anything
        # else here orphans every memory already written.
        assert runner.app_name == os.environ["AGENT_ENGINE_ID"]
    finally:
        for k, v in saved.items():
            if v is None:
                os.environ.pop(k, None)
            else:
                os.environ[k] = v


def test_a2a_extra_is_installed():
    # main_a2a.py imports this. Without the [a2a] extra it raises
    # ModuleNotFoundError, and the container dies on startup rather than at
    # any point a test would otherwise notice.
    import google.adk.a2a.utils.agent_to_a2a  # noqa: F401


def test_a2a_http_server_deps_are_installed():
    # Importing A2AStarletteApplication is NOT enough to prove the server can
    # run: `google-adk[a2a]` pulls a2a-sdk without its http-server extra, and
    # the missing sse-starlette only surfaces when the class is CONSTRUCTED,
    # inside to_a2a's lifespan startup. The class import passed, the suite was
    # green, and Cloud Run still died with "failed to start and listen on port
    # 8080". Assert the dependency itself.
    import sse_starlette  # noqa: F401
    import starlette  # noqa: F401


def test_agent_card_is_schema_valid_and_names_an_endpoint():
    from a2a.types import AgentCard

    # Imported from card.py, never from main_a2a: importing main_a2a would
    # call build_runner() and take the whole offline suite down.
    from Job_Helper_agent.card import CARD_PATH, load_agent_card

    # Parses as a real AgentCard, not just as JSON. A card that fails schema
    # validation takes the whole server down at import.
    raw = json.loads(CARD_PATH.read_text())
    AgentCard(**raw)

    # The url is what Gemini Enterprise calls back on. Passing a static card to
    # to_a2a() means ADK never fills this in (agent_to_a2a.py:203-205), so an
    # unresolved url is a silently unreachable agent.
    card = load_agent_card("job-helper-a2a-xyz.a.run.app", "https")
    assert card.url == "https://job-helper-a2a-xyz.a.run.app/"


def test_agent_card_declares_a2ui_and_streaming():
    card = json.loads(
        (pathlib.Path(__file__).parent / "Job_Helper_agent" / "agent_card.json").read_text()
    )
    caps = card["capabilities"]
    # Without streaming, progress updates cannot be painted incrementally.
    assert caps["streaming"] is True

    exts = caps["extensions"]
    a2ui = [e for e in exts if "a2ui" in e["uri"]]
    assert len(a2ui) == 1, "agent card must declare exactly one A2UI extension"

    # Gemini Enterprise supports A2UI v0.8 only. A version bump here silently
    # stops rendering in GE.
    assert a2ui[0]["uri"] == "https://a2ui.org/a2a-extension/a2ui/v0.8"

    # Must stay false: it is what lets the agent fall back to plain text for
    # any client that does not negotiate A2UI.
    assert a2ui[0]["required"] is False


def test_extract_user_id_reads_every_candidate_header_case_insensitively():
    # WHICH header Gemini Enterprise actually sends is UNCONFIRMED and must be
    # verified against a live GE call. Until then the lookup is deliberately
    # tolerant of several plausible names, and refuses rather than guessing
    # when none of them match -- see the module docstring in identity.py.
    from Job_Helper_agent.identity import IDENTITY_HEADERS, extract_user_id

    for header in IDENTITY_HEADERS:
        assert extract_user_id({header: "student@example.edu"}) == "student@example.edu"
        assert extract_user_id({header.upper(): "student@example.edu"}) == (
            "student@example.edu"
        )
        assert extract_user_id({header.title(): "student@example.edu"}) == (
            "student@example.edu"
        )


def test_extract_user_id_strips_the_google_issuer_prefix():
    # Google identity headers carry the issuer inline. Left on, it becomes part
    # of the user_id and silently forks every session and memory scope.
    from Job_Helper_agent.identity import extract_user_id

    assert (
        extract_user_id(
            {"x-goog-authenticated-user-email": "accounts.google.com:student@example.edu"}
        )
        == "student@example.edu"
    )


def test_extract_user_id_returns_none_rather_than_guessing():
    from Job_Helper_agent.identity import extract_user_id

    assert extract_user_id({}) is None
    assert extract_user_id({"x-goog-authenticated-user-email": ""}) is None
    assert extract_user_id({"x-goog-authenticated-user-email": "   "}) is None
    assert extract_user_id({"x-goog-authenticated-user-email": "accounts.google.com:"}) is None
    assert extract_user_id({"user-agent": "curl/8.0", "host": "example.com"}) is None


def test_request_converter_only_overrides_user_id_on_a_real_find():
    # No identity in the headers must leave the A2A_USER_* sentinel in place so
    # require_real_user still refuses the turn. Falling back to anything else
    # would hand one student's session to whoever calls next.
    from Job_Helper_agent.callbacks import require_real_user
    from Job_Helper_agent.identity import build_request_converter

    convert = build_request_converter()

    class _FakeRunRequest:
        user_id = "A2A_USER_ctx-abc123"

    class _FakeCallContext:
        def __init__(self, headers):
            self.state = {"headers": headers}

    class _FakeRequest:
        def __init__(self, headers):
            self.call_context = _FakeCallContext(headers)

    import Job_Helper_agent.identity as identity_module

    original = identity_module.convert_a2a_request_to_agent_run_request
    try:
        identity_module.convert_a2a_request_to_agent_run_request = (
            lambda request, part_converter: _FakeRunRequest()
        )

        found = convert(_FakeRequest({"x-goog-authenticated-user-email": "student@example.edu"}))
        assert found.user_id == "student@example.edu"
        assert require_real_user(_FakeCallbackContext("student@example.edu")) is None

        # No identity header: the sentinel must survive UNTOUCHED. The converter
        # must never invent an identity or reach for the service account.
        missing = convert(_FakeRequest({"user-agent": "curl/8.0"}))
        assert missing.user_id == "A2A_USER_ctx-abc123"
        # That sentinel is now allowed through, because it scopes to one
        # conversation rather than to one shared bucket -- see the guard test.
        assert require_real_user(_FakeCallbackContext(missing.user_id)) is None
    finally:
        identity_module.convert_a2a_request_to_agent_run_request = original


def test_require_public_host_refuses_the_localhost_default_on_cloud_run():
    # A deploy that forgets PUBLIC_HOST otherwise boots green while serving a
    # card pointing at https://localhost:8080/ -- a dead agent that looks
    # healthy. Cloud Run always sets K_SERVICE.
    from Job_Helper_agent.card import LOCAL_HOST_DEFAULT, require_public_host

    raised = None
    try:
        require_public_host(None, "job-helper-a2a")
    except RuntimeError as e:
        raised = e
    assert raised is not None, "boot succeeded on Cloud Run with no PUBLIC_HOST"
    assert "PUBLIC_HOST" in str(raised)

    assert (
        require_public_host("job-helper-a2a-xyz.a.run.app", "job-helper-a2a")
        == "job-helper-a2a-xyz.a.run.app"
    )

    # Local development keeps working with no environment at all.
    assert require_public_host(None, None) == LOCAL_HOST_DEFAULT


def test_specialists_emit_structured_data_for_rendering():
    """Cards need structure; prose cannot be drawn.

    company_agent and alumni_agent keep google_search -- ADK enforces the
    schema only on the final output, so the built-in still runs during the
    thought loop and the one-built-in-per-agent rule is untouched.
    """
    from Job_Helper_agent.agent import (
        alumni_structure_agent,
        company_structure_agent,
    )

    for agent in (company_structure_agent, company_structure_agent):
        assert agent.output_schema is not None, f"{agent.name} has no output_schema"
        # Rule 5: plain dict schemas, no Pydantic.
        assert isinstance(agent.output_schema, dict), f"{agent.name} schema is not a dict"

    assert alumni_structure_agent.output_schema is not None
    alumni_props = alumni_structure_agent.output_schema["properties"]["alumni"]["items"]
    # NO_INVENTION, structurally: a person without a found link is not a person
    # we may name, so the schema itself refuses to describe one.
    assert "profile_url" in alumni_props["required"]
    assert "name" in alumni_props["required"]


def test_company_and_alumni_surfaces_render_from_structured_state():
    from Job_Helper_agent.a2ui import build_a2ui_messages

    state = {
        "companies_data": {
            "companies": [
                {"name": "Zerodha", "why_it_fits": "Python + backend", "fit": "Strong",
                 "industry": "Fintech", "size_or_stage": "Mid-size",
                 "opening_title": "Backend Intern",
                 "opening_url": "https://zerodha.com/careers/backend-intern"},
            ]
        },
        "alumni_data": {
            "alumni": [
                {"name": "A. Rao", "role": "SDE", "company": "Zerodha",
                 "profile_url": "https://example.com/a-rao", "shared_context": "NITW 2019"},
            ],
            "search_links": {
                "people_search": "https://www.linkedin.com/search/results/people/?keywords=NITW+Zerodha",
                "school_page_search": "https://www.linkedin.com/search/results/schools/?keywords=NITW",
            },
        },
    }
    blob = json.dumps(build_a2ui_messages(state))
    assert "Zerodha" in blob and "Backend Intern" in blob
    assert "A. Rao" in blob and "https://example.com/a-rao" in blob


def test_alumni_surface_falls_back_to_linkedin_search_links():
    """No alumni found is the normal case, not an error.

    Public search cannot answer "who from college X works at company Y", so the
    student gets the LinkedIn searches to run themselves rather than nothing --
    and never a padded list of guessed names.
    """
    from Job_Helper_agent.a2ui import build_a2ui_messages

    state = {
        "alumni_data": {
            "alumni": [],
            "search_links": {
                "people_search": "https://www.linkedin.com/search/results/people/?keywords=NITW+Stripe",
                "school_page_search": "https://www.linkedin.com/search/results/schools/?keywords=NITW",
            },
        }
    }
    blob = json.dumps(build_a2ui_messages(state))
    assert "linkedin.com/search/results/people" in blob
    # Nothing may be invented to fill the gap.
    assert "A. Rao" not in blob

    # No links and no alumni -> that surface draws nothing at all.
    from Job_Helper_agent.a2ui import alumni_cards

    assert alumni_cards({"alumni": []}) is None
    assert alumni_cards({}) is None


def test_upload_prompt_shows_only_before_a_profile_exists():
    from Job_Helper_agent.a2ui import build_a2ui_messages

    empty = json.dumps(build_a2ui_messages({}))
    assert "resume" in empty.lower(), "no resume prompt on a cold start"

    warm = json.dumps(build_a2ui_messages({"profile": "B.Tech CSE, Python, NITW"}))
    assert "resume" not in warm.lower(), "still nagging after the profile exists"


def test_build_a2ui_messages_accepts_adk_state_not_just_dict():
    """ADK hands callbacks a `State`, and `dict(State)` raises `KeyError: 0`.

    dict() finds no keys() and falls back to reading it as a sequence of pairs.
    The caller swallows renderer errors so the student never loses their answer
    to a broken widget -- which meant this failed completely silently in
    production and the card simply never appeared.
    """
    from Job_Helper_agent.a2ui import build_a2ui_messages

    class _StateLike:
        def __init__(self, data):
            self._data = data

        def get(self, key, default=None):
            return self._data.get(key, default)

        def __getitem__(self, key):
            return self._data[key]

    raised = None
    try:
        dict(_StateLike({"applications": []}))
    except KeyError as e:
        raised = e
    assert raised is not None, "stand-in no longer reproduces the ADK State shape"

    state = _StateLike(
        {"applications": [{"company": "Zoho", "role": "Intern", "status": "Applied", "notes": ""}]}
    )
    msgs = build_a2ui_messages(state)
    assert msgs, "a State with applications must still render a board"
    assert "Zoho" in json.dumps(msgs)

    assert build_a2ui_messages(None) == []


def test_component_ids_are_namespaced_and_unique():
    """Ids must not collide with reserved words or across surfaces.

    Gemini Enterprise refused to draw a card with a component id of `body`
    ("This content could not be displayed. se`body`", observed 2026-07-28).
    `body` is also a valid Text usageHint value, and the official v0.8 fixture
    names that node `main-column`, never `body`. Several surfaces can render in
    one turn, so a bare `root` in each would collide too.
    """
    from Job_Helper_agent.a2ui import build_a2ui_messages

    state = {
        "applications": [{"company": "Zoho", "role": "Intern", "status": "Applied", "notes": ""}],
        "companies_data": {"companies": [{"name": "CRED", "why_it_fits": "python", "fit": "Strong"}]},
        "alumni_data": {"alumni": [], "search_links": {"people_search": "https://x/y"}},
    }
    msgs = build_a2ui_messages(state)
    surfaces = [m["surfaceUpdate"] for m in msgs if "surfaceUpdate" in m]
    assert len(surfaces) >= 3, "expected several surfaces in one turn"

    seen = set()
    reserved = {"body", "root", "head", "html", "title", "main"}
    for surface in surfaces:
        for component in surface["components"]:
            cid = component["id"]
            assert cid not in reserved, f"{cid!r} is a reserved/colliding id"
            assert cid not in seen, f"duplicate component id {cid!r} across surfaces"
            seen.add(cid)

    # Every beginRendering root must name a component that exists.
    roots = [m["beginRendering"]["root"] for m in msgs if "beginRendering" in m]
    assert len(roots) == len(surfaces)
    for root in roots:
        assert root in seen, f"beginRendering names missing component {root!r}"


def test_pipeline_board_renders_only_real_applications():
    import json

    from Job_Helper_agent.a2ui import (
        A2UI_MIME_TYPE,
        build_a2ui_messages,
        to_genai_parts,
    )

    # Nothing tracked -> no board. A student who has tracked nothing must not
    # get an empty board presented as if it were real. Asserted against the
    # builder directly: build_a2ui_messages also draws the upload prompt on a
    # cold start, which is a different surface with its own test.
    from Job_Helper_agent.a2ui import pipeline_board

    assert pipeline_board([]) is None
    assert pipeline_board([{"company": "", "role": "x", "status": "Applied"}]) is None

    apps = [
        {"company": "Freshworks", "role": "SWE Intern", "status": "Applied", "notes": ""},
        {"company": "Zoho", "role": "Backend Intern", "status": "Interview", "notes": "onsite Aug 3"},
    ]
    # A profile is present so the cold-start upload prompt does not also draw;
    # this test is about the board alone.
    msgs = build_a2ui_messages({"profile": "B.Tech CSE", "applications": apps})

    # Flat list of messages, surfaceUpdate before beginRendering.
    kinds = [next(iter(m)) for m in msgs]
    assert kinds == ["surfaceUpdate", "beginRendering"], kinds

    comps = msgs[0]["surfaceUpdate"]["components"]
    ids = {c["id"] for c in comps}
    root = msgs[1]["beginRendering"]["root"]
    assert root in ids

    # Every referenced child must exist, or the renderer drops the subtree.
    for c in comps:
        body = next(iter(c["component"].values()))
        refs = []
        if "child" in body:
            refs.append(body["child"])
        refs += (body.get("children") or {}).get("explicitList", [])
        for r in refs:
            assert r in ids, f"{c['id']} references missing component {r}"

    blob = json.dumps(msgs)
    assert "Freshworks" in blob and "Interview" in blob
    # NO_INVENTION, structurally: nothing may appear that was not in state.
    assert "Google" not in blob

    parts = to_genai_parts(msgs)
    assert len(parts) == len(msgs)
    for p in parts:
        raw = p.inline_data.data.decode("utf-8")
        assert raw.startswith("<a2a_datapart_json>")
        inner = json.loads(raw[len("<a2a_datapart_json>") : -len("</a2a_datapart_json>")])
        # Gemini Enterprise identifies an A2UI payload by this mime on the
        # DataPart metadata. Wrong mime renders as raw JSON text.
        assert inner["metadata"]["mimeType"] == A2UI_MIME_TYPE


# ---------------------------------------------------------------------------
# placement_agent -- reading the file the user uploaded in Gemini Enterprise
#
# GE never gives a custom agent a file path. It announces the upload as a text
# marker ("<start_of_user_uploaded_file: resume.pdf>", empty between markers)
# and puts the bytes in the artifact service. Measured against a live deployed
# agent 2026-07-22; see .claude/skills/gemini-enterprise-agents. A parser that
# only stats the filesystem can never succeed in the deployed container, and
# the model then invents the resume contents instead of erroring.
# ---------------------------------------------------------------------------

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class _FakeUploadContext:
    """Stand-in for ToolContext: only the artifact methods are touched."""

    def __init__(self, artifacts=None, has_artifact_service=True):
        self._artifacts = dict(artifacts or {})
        self._has_service = has_artifact_service
        self.state = {}

    async def list_artifacts(self):
        if not self._has_service:
            raise ValueError("Artifact service is not initialized.")
        return list(self._artifacts)

    async def load_artifact(self, filename, version=None):
        if not self._has_service:
            raise ValueError("Artifact service is not initialized.")
        return self._artifacts.get(filename)


def _docx_bytes(paragraphs) -> bytes:
    from docx import Document

    doc = Document()
    for line in paragraphs:
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _upload(data: bytes, mime_type: str):
    return types.Part.from_bytes(data=data, mime_type=mime_type)


def _pdf_bytes(line: str) -> bytes:
    """A minimal one-page PDF with a real text layer -- no extra dependency."""
    stream = f"BT /F1 12 Tf 72 720 Td ({line}) Tj ET".encode()
    objects = [
        b"<</Type/Catalog/Pages 2 0 R>>",
        b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]"
        b"/Resources<</Font<</F1 5 0 R>>>>/Contents 4 0 R>>",
        b"<</Length " + str(len(stream)).encode() + b">>\nstream\n" + stream + b"\nendstream\n",
        b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    ]
    buf, offsets = io.BytesIO(), []
    buf.write(b"%PDF-1.4\n")
    for i, obj in enumerate(objects, 1):
        offsets.append(buf.tell())
        buf.write(str(i).encode() + b" 0 obj\n" + obj + b"\nendobj\n")
    xref = buf.tell()
    buf.write(b"xref\n0 " + str(len(objects) + 1).encode() + b"\n0000000000 65535 f \n")
    for off in offsets:
        buf.write(b"%010d 00000 n \n" % off)
    buf.write(
        b"trailer<</Size " + str(len(objects) + 1).encode() + b"/Root 1 0 R>>\nstartxref\n"
        + str(xref).encode() + b"\n%%EOF\n"
    )
    return buf.getvalue()


def test_parse_resume_reads_an_uploaded_pdf_artifact():
    """PDF is the format users actually upload; extraction runs off bytes."""
    ctx = _FakeUploadContext(
        {"resume.pdf": _upload(_pdf_bytes("Priya Raman Backend Intern Python"), "application/pdf")}
    )
    out = asyncio.run(parse_resume(ctx, "resume.pdf"))

    assert out["success"], out
    assert "Priya Raman" in out["text"]
    assert out["source"] == "uploaded_file"


def test_parse_resume_reads_an_uploaded_docx_artifact():
    """The deployed path: bytes come from the artifact service, not from disk."""
    ctx = _FakeUploadContext(
        {"resume.docx": _upload(_docx_bytes(["Priya Raman", "Backend Intern"]), DOCX_MIME)}
    )
    out = asyncio.run(parse_resume(ctx, "resume.docx"))

    assert out["success"], out
    assert "Priya Raman" in out["text"]
    assert out["source"] == "uploaded_file"
    assert out["word_count"] > 0


def test_parse_resume_reads_docx_table_cells():
    """Resumes are routinely laid out in tables -- paragraphs alone drop them."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Priya Raman")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "Python, Kubernetes"
    buf = io.BytesIO()
    doc.save(buf)

    ctx = _FakeUploadContext({"resume.docx": _upload(buf.getvalue(), DOCX_MIME)})
    out = asyncio.run(parse_resume(ctx, "resume.docx"))

    assert out["success"], out
    assert "Kubernetes" in out["text"], "table cell content was silently dropped"


def test_parse_resume_accepts_the_ge_upload_marker_verbatim():
    """The model sees the marker, not a bare name -- both must resolve."""
    ctx = _FakeUploadContext({"resume.txt": _upload(b"Priya Raman\nBackend Intern", "text/plain")})
    out = asyncio.run(parse_resume(ctx, "<start_of_user_uploaded_file: resume.txt>"))

    assert out["success"], out
    assert "Priya Raman" in out["text"]


def test_parse_resume_finds_the_only_upload_without_being_told_its_name():
    ctx = _FakeUploadContext({"my cv.txt": _upload(b"Priya Raman", "text/plain")})
    out = asyncio.run(parse_resume(ctx))

    assert out["success"], out
    assert out["file_name"] == "my cv.txt"


def test_parse_resume_reports_what_is_attached_when_the_name_is_wrong():
    ctx = _FakeUploadContext({"resume.txt": _upload(b"Priya Raman", "text/plain")})
    out = asyncio.run(parse_resume(ctx, "not-the-file.pdf"))

    assert out["success"] is False
    assert out["available_files"] == ["resume.txt"]


def test_parse_resume_says_nothing_was_uploaded_rather_than_returning_blank():
    """A blank success would let the model hallucinate a resume."""
    ctx = _FakeUploadContext({})
    out = asyncio.run(parse_resume(ctx))

    assert out["success"] is False
    assert out["text"] == ""
    assert "upload" in out["error"].lower()


def test_parse_resume_still_reads_a_local_path_for_adk_web():
    """Local `adk web` runs have no artifact service; a path must still work."""
    fd, path = tempfile.mkstemp(suffix=".txt")
    with os.fdopen(fd, "w", encoding="utf-8") as fh:
        fh.write("Priya Raman\nBackend Intern")
    try:
        out = asyncio.run(parse_resume(_FakeUploadContext(has_artifact_service=False), path))
    finally:
        os.unlink(path)

    assert out["success"], out
    assert out["source"] == "local_path"
    assert "Priya Raman" in out["text"]


def test_parse_resume_flags_a_pdf_with_no_text_layer():
    """A scanned resume must be reported, not returned as an empty success."""
    empty_pdf = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]>>endobj\n"
        b"trailer<</Root 1 0 R>>\n%%EOF\n"
    )
    ctx = _FakeUploadContext({"scan.pdf": _upload(empty_pdf, "application/pdf")})
    out = asyncio.run(parse_resume(ctx, "scan.pdf"))

    assert out["success"] is False
    assert out["text"] == ""


def test_placement_root_is_told_the_upload_marker_protocol():
    """Guardrail: lose the marker instruction and the model invents resumes."""
    instruction = placement_root_agent.instruction
    assert "<start_of_user_uploaded_file:" in instruction
    assert "parse_resume" in instruction


def test_placement_agents_do_not_mix_builtin_and_function_tools():
    agents = [placement_root_agent, *(placement_root_agent.sub_agents or [])]
    for agent in agents:
        tools = list(agent.tools or [])
        built_ins = [t for t in tools if _is_built_in(t)]
        if built_ins:
            assert len(tools) == 1, (
                f"{agent.name} holds built-in {_tool_name(built_ins[0])!r} alongside"
                " function tools. Gemini rejects this at request time."
            )


if __name__ == "__main__":
    for name, fn in sorted(globals().items()):
        if name.startswith("test_") and callable(fn):
            fn()
            print(f"ok  {name}")
    print("\nall checks passed")
