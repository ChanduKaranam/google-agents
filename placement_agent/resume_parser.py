"""
resume_parser.py
Reads the resume the user attached to the chat.

Gemini Enterprise does NOT hand a custom agent a file path. It announces the
upload as a pair of text markers with nothing between them --
"<start_of_user_uploaded_file: resume.pdf>" -- and puts the bytes in the ADK
artifact service. So the file is fetched with tool_context.load_artifact(), and
a filesystem path is only the fallback for local `adk web` runs.
(Measured against a live deployed agent 2026-07-22; see the
gemini-enterprise-agents skill.)
"""

from __future__ import annotations

import io
import os
import re
from typing import Optional

from google.adk.tools.tool_context import ToolContext

# The model usually echoes the whole marker back as the filename.
_MARKER = re.compile(r"<(?:start|end)_of_user_uploaded_file:\s*(?P<name>[^>]+)>")


# ---------------------------------------------------------------------------
# Byte-level extraction
# ---------------------------------------------------------------------------

def _pdf_text(data: bytes) -> str:
    try:
        import pdfplumber
    except ImportError:
        return "ERROR: pdfplumber is not installed. Run: pip install pdfplumber"

    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = (page.extract_text() for page in pdf.pages)
            return "\n".join(p for p in pages if p).strip()
    except Exception as e:
        return f"ERROR reading PDF: {e}"


def _docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        return "ERROR: python-docx is not installed. Run: pip install python-docx"

    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        return f"ERROR reading DOCX: {e}"

    chunks = [p.text for p in doc.paragraphs if p.text.strip()]
    # Two-column resume templates put everything in tables; doc.paragraphs
    # silently returns none of it.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    chunks.append(cell.text.strip())
    return "\n".join(chunks).strip()


def _extract_text(data: bytes, file_name: str, mime_type: str = "") -> str:
    """Turn resume bytes into plain text, sniffing the format rather than
    trusting the extension (GE filenames are not always faithful)."""
    ext = os.path.splitext(file_name)[1].lower()
    mime = (mime_type or "").lower()

    if data[:5] == b"%PDF-" or ext == ".pdf" or "pdf" in mime:
        return _pdf_text(data)
    if data[:2] == b"PK" or ext in (".docx", ".doc") or "wordprocessingml" in mime:
        return _docx_text(data)
    if ext in (".txt", ".md", ".rtf") or mime.startswith("text/"):
        return data.decode("utf-8", errors="replace").strip()

    # Unknown label: accept it if it decodes as readable text, else refuse.
    text = data.decode("utf-8", errors="ignore").strip()
    if text and sum(c.isprintable() or c.isspace() for c in text) > 0.9 * len(text):
        return text
    return (
        f"ERROR: Unsupported file type '{ext or mime or 'unknown'}'. "
        "Please upload the resume as a PDF, DOCX, or plain-text file."
    )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _clean_name(raw: str) -> str:
    marker = _MARKER.search(raw or "")
    if marker:
        return marker.group("name").strip()
    return (raw or "").strip().strip("\"'")


async def _uploaded_files(tool_context: ToolContext) -> list[str]:
    """Artifacts attached to this session, or [] when there is no artifact
    service (local runs, and any direct unit call)."""
    try:
        return list(await tool_context.list_artifacts() or [])
    except (ValueError, AttributeError):
        return []


def _fail(error: str, file_name: str = "", available: Optional[list] = None) -> dict:
    return {
        "success": False,
        "text": "",
        "error": error,
        "file_name": os.path.basename(file_name),
        "word_count": 0,
        "source": "",
        "available_files": list(available or []),
    }


# ---------------------------------------------------------------------------
# Public ADK tool functions
# ---------------------------------------------------------------------------

async def parse_resume(tool_context: ToolContext, file_name: str = "") -> dict:
    """
    Read the resume the user attached to this chat and return its plain text.

    Args:
        file_name: Name of the attached file, as it appears in the
                   "<start_of_user_uploaded_file: NAME>" marker. Leave empty to
                   use the only file attached to the conversation. A local file
                   path also works when running outside Gemini Enterprise.

    Returns:
        A dict with keys:
          - success (bool)
          - text (str): the extracted resume text on success
          - error (str): what went wrong, and what to ask the user for
          - file_name (str): the file actually read
          - word_count (int)
          - source (str): "uploaded_file" or "local_path"
          - available_files (list): files attached to this chat, on failure
    """
    wanted = _clean_name(file_name)
    uploaded = await _uploaded_files(tool_context)

    if not wanted:
        if len(uploaded) == 1:
            wanted = uploaded[0]
        elif uploaded:
            return _fail(
                "Several files are attached to this chat. Ask the user which one "
                "is their resume, then call parse_resume again with that name.",
                available=uploaded,
            )
        else:
            return _fail(
                "No file has been uploaded in this chat yet. Ask the user to attach "
                "their resume (PDF or DOCX) with the upload button and send it again. "
                "Do not analyse a resume you have not read.",
            )

    # GE occasionally reports a different case or a path-prefixed name.
    if wanted not in uploaded:
        base = os.path.basename(wanted).lower()
        wanted = next(
            (n for n in uploaded if n.lower() == wanted.lower() or os.path.basename(n).lower() == base),
            wanted,
        )

    if wanted in uploaded:
        source = "uploaded_file"
        part = await tool_context.load_artifact(wanted)
        if part is None:
            return _fail(
                f"'{wanted}' is attached but could not be loaded. Ask the user to "
                "re-upload it.",
                wanted,
                uploaded,
            )
        blob = getattr(part, "inline_data", None)
        if blob is None:
            text = (getattr(part, "text", "") or "").strip()
        else:
            text = _extract_text(blob.data or b"", wanted, blob.mime_type or "")
    elif os.path.exists(wanted):
        # Local `adk web` / CLI runs, where a real path is meaningful.
        source = "local_path"
        with open(wanted, "rb") as fh:
            text = _extract_text(fh.read(), wanted)
    else:
        return _fail(
            f"No file named '{wanted}' is attached to this chat. Ask the user to "
            "upload their resume again.",
            wanted,
            uploaded,
        )

    if text.startswith("ERROR"):
        return _fail(text, wanted, uploaded)

    if not text:
        return _fail(
            f"'{os.path.basename(wanted)}' contains no readable text -- it is most "
            "likely a scanned or image-only PDF. Ask the user for a text-based PDF "
            "or DOCX export. Do not guess at its contents.",
            wanted,
            uploaded,
        )

    return {
        "success": True,
        "text": text,
        "error": "",
        "file_name": os.path.basename(wanted),
        "word_count": len(text.split()),
        "source": source,
        "available_files": uploaded,
    }


async def list_uploaded_files(tool_context: ToolContext) -> dict:
    """
    List the files the user has attached to this conversation.

    Use this when the user says they uploaded something but parse_resume cannot
    find it, so you can tell them exactly what did and did not arrive.

    Returns:
        A dict with 'files' (list of filenames) and 'total' (int).
    """
    files = await _uploaded_files(tool_context)
    return {"files": files, "total": len(files)}


def get_resume_sections(resume_text: str) -> dict:
    """
    Identify and extract standard resume sections from plain text.

    Args:
        resume_text: The raw text content of a resume.

    Returns:
        A dict mapping section names to their extracted content.
        Sections attempted: contact, summary, experience, education,
        skills, certifications, projects, awards.
    """
    # Common section header patterns (case-insensitive)
    section_patterns = {
        "contact": r"(contact\s*(information|details)?)",
        "summary": r"(summary|objective|profile|about\s*me)",
        "experience": r"(experience|work\s*(history|experience)|employment)",
        "education": r"(education|academic|qualification)",
        "skills": r"(skills|technical\s*skills|core\s*competencies|competencies)",
        "certifications": r"(certifications?|licenses?|credentials?)",
        "projects": r"(projects?|portfolio)",
        "awards": r"(awards?|honors?|achievements?|accomplishments?)",
    }

    lines = resume_text.splitlines()
    found_sections: dict[str, str] = {}
    current_section: Optional[str] = None
    buffer: list[str] = []

    for line in lines:
        stripped = line.strip()
        matched_section = None
        for section_name, pattern in section_patterns.items():
            if re.fullmatch(pattern, stripped, re.IGNORECASE):
                matched_section = section_name
                break
        if matched_section:
            # Save previous buffer
            if current_section and buffer:
                found_sections[current_section] = "\n".join(buffer).strip()
            current_section = matched_section
            buffer = []
        elif current_section:
            buffer.append(line)

    # Save last section
    if current_section and buffer:
        found_sections[current_section] = "\n".join(buffer).strip()

    # If no sections were detected at all, return the full text as "body"
    if not found_sections:
        found_sections["body"] = resume_text

    return {
        "sections_found": list(found_sections.keys()),
        "sections": found_sections,
        "total_sections": len(found_sections),
    }
